import os
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
import requests
from bs4 import BeautifulSoup
import pypdf
import docx

# Initialize structured logging
logger = logging.getLogger("rag_system.ingestion.document_loader")


class DocumentLoadError(Exception):
    """Raised when document loading fails."""
    pass


class Document:
    """Represents a parsed document with its textual content and metadata."""

    def __init__(self, text: str, metadata: Dict[str, Any]) -> None:
        """Initializes a Document instance.

        Args:
            text: The textual content of the document.
            metadata: Metadata mapping (e.g. source, title, page_number, loaded_at).
        """
        self.text = text
        self.metadata = metadata

    def __repr__(self) -> str:
        return f"Document(title={self.metadata.get('title')}, page={self.metadata.get('page_number')}, source={self.metadata.get('source')})"


class DocumentLoader:
    """Service to load and parse documents from various formats (PDF, DOCX, MD, HTML)."""

    def __init__(self) -> None:
        """Initializes the DocumentLoader."""
        pass

    def load_pdf(self, file_path: str) -> List[Document]:
        """Loads and parses a PDF file using pypdf.

        Args:
            file_path: Path to the PDF file on the local filesystem.

        Returns:
            A list of Document objects (one per page).

        Raises:
            DocumentLoadError: If the PDF is corrupt, missing, or unparseable.
        """
        logger.info(f"Loading PDF file: {file_path}")
        if not os.path.exists(file_path):
            raise DocumentLoadError(f"PDF file not found: {file_path}")

        documents: List[Document] = []
        try:
            reader = pypdf.PdfReader(file_path)
            title = os.path.basename(file_path)
            if reader.metadata and reader.metadata.title:
                title = reader.metadata.title

            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                metadata = {
                    "source": os.path.abspath(file_path),
                    "title": title,
                    "page_number": page_idx + 1,
                    "loaded_at": datetime.utcnow().isoformat()
                }
                documents.append(Document(text=text, metadata=metadata))
        except Exception as e:
            error_msg = f"Failed to parse PDF {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentLoadError(error_msg) from e

        logger.info(f"Successfully loaded {len(documents)} pages from PDF: {file_path}")
        return documents

    def load_markdown(self, file_path: str) -> List[Document]:
        """Loads and parses a markdown file.

        Args:
            file_path: Path to the markdown file on the local filesystem.

        Returns:
            A list containing a single Document object representing the file.

        Raises:
            DocumentLoadError: If the file is unreadable.
        """
        logger.info(f"Loading markdown file: {file_path}")
        if not os.path.exists(file_path):
            raise DocumentLoadError(f"Markdown file not found: {file_path}")

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            title = os.path.basename(file_path)
            # Try to extract the first # Header as the title
            for line in text.splitlines():
                if line.strip().startswith("#"):
                    title = line.strip().lstrip("#").strip()
                    break

            metadata = {
                "source": os.path.abspath(file_path),
                "title": title,
                "page_number": 1,
                "loaded_at": datetime.utcnow().isoformat()
            }
            return [Document(text=text, metadata=metadata)]
        except Exception as e:
            error_msg = f"Failed to read markdown file {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentLoadError(error_msg) from e

    def load_docx(self, file_path: str) -> List[Document]:
        """Loads and parses a Word Document (.docx).

        Args:
            file_path: Path to the docx file on the local filesystem.

        Returns:
            A list containing a single Document object representing the document.

        Raises:
            DocumentLoadError: If the docx is corrupt or missing.
        """
        logger.info(f"Loading DOCX file: {file_path}")
        if not os.path.exists(file_path):
            raise DocumentLoadError(f"docx file not found: {file_path}")

        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = "\n".join(full_text)

            title = os.path.basename(file_path)
            metadata = {
                "source": os.path.abspath(file_path),
                "title": title,
                "page_number": 1,
                "loaded_at": datetime.utcnow().isoformat()
            }
            return [Document(text=text, metadata=metadata)]
        except Exception as e:
            error_msg = f"Failed to read DOCX file {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentLoadError(error_msg) from e

    def load_web_page(self, url: str) -> List[Document]:
        """Fetches and parses a web page using requests and beautifulsoup4.

        Args:
            url: HTTP/HTTPS URL of the web page.

        Returns:
            A list containing a single Document object representing the web page.

        Raises:
            DocumentLoadError: If the network request fails or HTML is invalid.
        """
        logger.info(f"Fetching web page: {url}")
        try:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
        except Exception as e:
            error_msg = f"Network request failed for URL {url}: {str(e)}"
            logger.error(error_msg)
            raise DocumentLoadError(error_msg) from e

        try:
            soup = BeautifulSoup(response.text, "html.parser")
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text(separator="\n")
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = "\n".join(chunk for chunk in chunks if chunk)

            title = soup.title.string.strip() if soup.title else url
            metadata = {
                "source": url,
                "title": title,
                "page_number": 1,
                "loaded_at": datetime.utcnow().isoformat()
            }
            return [Document(text=clean_text, metadata=metadata)]
        except Exception as e:
            error_msg = f"Failed to parse HTML from URL {url}: {str(e)}"
            logger.error(error_msg)
            raise DocumentLoadError(error_msg) from e

    def load_directory(self, dir_path: str, recursive: bool = False) -> List[Document]:
        """Scans a directory for supported file formats and parses them.

        Supported file extensions: .pdf, .md, .docx, .txt

        Args:
            dir_path: Local filesystem directory path.
            recursive: Whether to scan subdirectories recursively.

        Returns:
            A list of all successfully parsed Document objects.

        Raises:
            DocumentLoadError: If the directory does not exist.
        """
        logger.info(f"Scanning directory: {dir_path} (recursive={recursive})")
        if not os.path.isdir(dir_path):
            raise DocumentLoadError(f"Directory not found: {dir_path}")

        documents: List[Document] = []
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                full_path = os.path.join(root, file)
                try:
                    if ext == ".pdf":
                        documents.extend(self.load_pdf(full_path))
                    elif ext == ".md":
                        documents.extend(self.load_markdown(full_path))
                    elif ext == ".docx":
                        documents.extend(self.load_docx(full_path))
                    elif ext == ".txt":
                        documents.extend(self.load_markdown(full_path))  # txt files can load as markdown/raw text
                except Exception as e:
                    logger.warning(f"Skipping file {full_path} due to loading error: {str(e)}")

            if not recursive:
                break

        logger.info(f"Total documents loaded from directory '{dir_path}': {len(documents)}")
        return documents
